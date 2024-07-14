import re
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert

last_period = None

def convert_time(match):
    global last_period
    time = match.group(0)
    print(f'Original time: {time}')
    hours, minutes, period = re.match(r'(\d{1,2}):(\d{2})\s*(AM|PM)?', time, re.I).groups()
    if period:
        last_period = period
    else:
        period = last_period
    hours = int(hours)
    if period and period.upper() == 'PM' and hours != 12:
        hours += 12
    elif period and period.upper() == 'AM' and hours == 12:
        hours = 0
    new_time = f'{hours:02d}:{minutes}'
    print(f'Converted time: {new_time}')
    return new_time

def convert_pdf_times(input_filename, output_filename):
    print('Converting PDF to Word...')
    # Convert the PDF to Word
    cv = Converter(input_filename)
    cv.convert('temp.docx', start=0, end=None)
    cv.close()

    print('Opening Word document and replacing times...')
    # Open the Word document and replace times
    doc = Document('temp.docx')
    for i, paragraph in enumerate(doc.paragraphs):
        print(f'Paragraph {i}: {paragraph.text}')
        if re.search(r'(\d{1,2}:\d{2})\s*(AM|PM)?', paragraph.text, re.I):
            for run in paragraph.runs:
                run.text = re.sub(r'(\d{1,2}:\d{2})\s*(AM|PM)?', convert_time, run.text, flags=re.I)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(f'Table cell: {cell.text}')
                if re.search(r'(\d{1,2}:\d{2})\s*(AM|PM)?', cell.text, re.I):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = re.sub(r'(\d{1,2}:\d{2})\s*(AM|PM)?', convert_time, run.text, flags=re.I)
    doc.save('temp.docx')

    print('Converting Word document back to PDF...')
    # Convert the Word document back to PDF
    convert('temp.docx', output_filename)

    print('Done.')

convert_pdf_times('schedule.pdf', 'new-schedule.pdf')